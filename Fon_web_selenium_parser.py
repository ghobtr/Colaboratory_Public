# coding=utf-8

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import pathlib
from docx import Document
from datetime import datetime, timedelta
import time
import os
from docx.shared import Inches
import pandas as pd

# Requirements
# pip install selenium
# pip install python-docx
# pip install openpyxl

# Date var
now = datetime.now()  # current date and time
past = datetime.now() - timedelta(7)  # for last week
past2 = datetime.now() - timedelta(14)  # for 2 last week
date_time = now.strftime("%d/%m/%Y, %H:%M")
date_time_file = now.strftime("%d_%m_%Y")
today = now.strftime("%d.%m.%Y")
past = past.strftime("%d.%m.%Y")
past2 = past2.strftime("%d.%m.%Y")


# selenium var
current_path = pathlib.Path().absolute()
chrome_options = Options()
chrome_options.add_argument('--headless')
chrome_options.add_argument('--start-maximized')
prefs = {"download.default_directory" : str(current_path),
		"download.prompt_for_download": False}
chrome_options.add_experimental_option('prefs', prefs)



# clear old files
try:
	os.remove(str(current_path) + "/FonTurkey_Fon_Karsilastirma.xlsx")
	print("old files are deleted")
except:
	pass




chromedriver = str(current_path) + "/chromedriver"
driver = webdriver.Chrome(executable_path=chromedriver, options=chrome_options)

fon_list = ["NNF"]
#fon_list = ["NNF","TCD","AOY","AFS","MAC","AFV","AFA","TGE","GTA","GPU","GUH","GPB","OPH","OPI","ODV","ARM"]

# begin program

document = Document()
document.add_heading('Fon Report', 0)
document.add_paragraph("Rapor üretilme zamanı: " + str(date_time))

for fon_name in fon_list:
	the_url = "https://www.tefas.gov.tr/FonAnaliz.aspx?FonKod=" + fon_name
	driver.get(the_url)
	# xpath 
	fon_name_xpath = '//span[@id="MainContent_FormViewMainIndicators_LabelFund"]'
	day_price_xpath = '//li[text()="Günlük Getiri (%)"]/span'
	last_month_price_xpath = '//li[text()="Son 1 Ay Getirisi"]/span'
	last_year_price_xpath = '//li[text()="Son 1 Yıl Getirisi"]/span'

	fon_detail_name = driver.find_element_by_xpath(fon_name_xpath)
	day_price = driver.find_element_by_xpath(day_price_xpath)
	last_month_price = driver.find_element_by_xpath(last_month_price_xpath)
	last_year_price = driver.find_element_by_xpath(last_year_price_xpath)

	# Generate docx
	document.add_heading(fon_name, level=1)
	document.add_paragraph(fon_detail_name.text)
	document.add_paragraph("Günlük Getiri")
	document.add_paragraph(day_price.text)
	document.add_paragraph("Aylık Getiri")
	document.add_paragraph(last_month_price.text)
	document.add_paragraph("Yıllık Getiri")
	document.add_paragraph(last_year_price.text)

# download xlsx
the_url = "https://www.tefas.gov.tr/FonKarsilastirma.aspx"
driver.get(the_url)
input_element_start_date=driver.find_element_by_css_selector('input[id = "MainContent_TextBoxStartDate"]')
input_element_start_date.send_keys(past)
input_element_finish_date=driver.find_element_by_css_selector('input[id = "MainContent_TextBoxEndDate"]')
input_element_finish_date.send_keys(today)
submit_button=driver.find_element_by_css_selector('input#MainContent_ButtonSearchDates.fund-compare-button')
submit_button.click()

time.sleep(2)
driver.set_window_size(1280, 1800)
time.sleep(2)
screent_shoot = str(current_path) + "/1.png"
driver.save_screenshot(screent_shoot)

save_button=driver.find_element_by_css_selector('input#MainContent_ImageButtonExcelGenel.export-img')
save_button.click()
time.sleep(2)

driver.execute_script("""__doPostBack('ctl00$MainContent$GridViewFundReturn','Page$2')""")
time.sleep(2)
driver.set_window_size(1280, 1800)
time.sleep(2)
screent_shoot = str(current_path) + "/1a.png"
driver.save_screenshot(screent_shoot)

## XLSX
xls = pd.ExcelFile(str(current_path) + "/FonTurkey_Fon_Karsilastirma.xlsx",engine='openpyxl')
df1 = pd.read_excel(xls, 'Getiri Bazlı')



df1 = df1[~df1['Şemsiye Fon Türü'].isin(['Serbest Şemsiye Fonu'])]

df1=df1.sort_values(by='Getiri', ascending=False)
df2=df1.sort_values(by='Getiri', ascending=True)



document.add_page_break()




document.add_heading("Geçen Haftanın en iyileri - Tablo ", level=1)
document.add_paragraph('Tarih aralığı ' + str(past) + "----" + str(today), style='Intense Quote')


table = document.add_table(rows=1, cols=4,style='Table Grid')
table.allow_autofit=True


hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fon Kodu'
hdr_cells[1].text = 'Fon Adı'
hdr_cells[2].text = 'Şemsiye Fon Türü'
hdr_cells[3].text = 'Getiri'

for i in range (20):
	row_cells = table.add_row().cells
	row_cells[0].text = df1.iloc[i]['Fon Kodu']
	row_cells[1].text = df1.iloc[i]['Fon Adı']
	row_cells[2].text = df1.iloc[i]['Şemsiye Fon Türü']
	row_cells[3].text = str(df1.iloc[i]['Getiri'])
###

document.add_page_break()
document.add_heading('Geçen Hafta Bizimkiler - Tablo' , level=1)
document.add_paragraph('Tarih aralığı ' + str(past) + "----" + str(today), style='Intense Quote')
table = document.add_table(rows=1, cols=4,style='Table Grid')
table.allow_autofit=True

parsing_df1=df1[df1['Fon Kodu'].str.contains("NNF|TCD|AOY|AFS|MAC|AFV|AFA|TGE|GTA|GPU|GUH|GPB|OPH|OPI|ODV|ARM") == True] 

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fon Kodu'
hdr_cells[1].text = 'Fon Adı'
hdr_cells[2].text = 'Şemsiye Fon Türü'
hdr_cells[3].text = 'Getiri'

for i in range (len(parsing_df1)):
	row_cells = table.add_row().cells
	row_cells[0].text = parsing_df1.iloc[i]['Fon Kodu']
	row_cells[1].text = parsing_df1.iloc[i]['Fon Adı']
	row_cells[2].text = parsing_df1.iloc[i]['Şemsiye Fon Türü']
	row_cells[3].text = str(parsing_df1.iloc[i]['Getiri'])


###
# Kaybedenler Geçen hafta
document.add_page_break()

document.add_heading("Geçen Haftanın en Kötüler - Tablo ", level=1)
document.add_paragraph('Tarih aralığı ' + str(past) + "----" + str(today), style='Intense Quote')


table = document.add_table(rows=1, cols=4,style='Table Grid')
table.allow_autofit=True


hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fon Kodu'
hdr_cells[1].text = 'Fon Adı'
hdr_cells[2].text = 'Şemsiye Fon Türü'
hdr_cells[3].text = 'Getiri'

for i in range (20):
	row_cells = table.add_row().cells
	row_cells[0].text = df2.iloc[i]['Fon Kodu']
	row_cells[1].text = df2.iloc[i]['Fon Adı']
	row_cells[2].text = df2.iloc[i]['Şemsiye Fon Türü']
	row_cells[3].text = str(df2.iloc[i]['Getiri'])
###



document.add_page_break()

###




document.add_page_break()
document.add_heading('Geçen Haftanın En İyileri - Resim' , level=1)
document.add_paragraph('Tarih aralığı ' + str(past) + "----" + str(today), style='Intense Quote')

p = document.add_paragraph()
r = p.add_run()
r.add_picture(str(current_path) + "/1.png",width=Inches(5.5), height=Inches(7.7))

r = p.add_run()
r.add_picture(str(current_path) + "/1a.png",width=Inches(5.5), height=Inches(7.7))

the_url = "https://www.tefas.gov.tr/FonKarsilastirma.aspx"
driver.get(the_url)
input_element_start_date=driver.find_element_by_css_selector('input[id = "MainContent_TextBoxStartDate"]')
input_element_start_date.send_keys(past2)
input_element_finish_date=driver.find_element_by_css_selector('input[id = "MainContent_TextBoxEndDate"]')
input_element_finish_date.send_keys(today)
submit_button=driver.find_element_by_css_selector('input#MainContent_ButtonSearchDates.fund-compare-button')
submit_button.click()



time.sleep(2)
driver.set_window_size(1280, 1800)
time.sleep(2)
screent_shoot = str(current_path) + "/2.png"
driver.save_screenshot(screent_shoot)

time.sleep(2)
os.remove(str(current_path) + "/FonTurkey_Fon_Karsilastirma.xlsx")


save_button=driver.find_element_by_css_selector('input#MainContent_ImageButtonExcelGenel.export-img')
save_button.click()
time.sleep(2)

driver.execute_script("""__doPostBack('ctl00$MainContent$GridViewFundReturn','Page$2')""")
time.sleep(2)
driver.set_window_size(1280, 1800)
time.sleep(2)
screent_shoot = str(current_path) + "/2a.png"
driver.save_screenshot(screent_shoot)

driver.close()

## XLSX
xls = pd.ExcelFile(str(current_path) + "/FonTurkey_Fon_Karsilastirma.xlsx",engine='openpyxl')
df1 = pd.read_excel(xls, 'Getiri Bazlı')



df1 = df1[~df1['Şemsiye Fon Türü'].isin(['Serbest Şemsiye Fonu'])]

df1=df1.sort_values(by='Getiri', ascending=False)
df2=df1.sort_values(by='Getiri', ascending=True)


document.add_page_break()

document.add_heading("Geçen 2 Haftanın en iyileri - Tablo", level=1)
document.add_paragraph('Tarih aralığı ' + str(past2) + "----" + str(today), style='Intense Quote')
table = document.add_table(rows=1, cols=4,style='Table Grid')
table.allow_autofit=True


hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fon Kodu'
hdr_cells[1].text = 'Fon Adı'
hdr_cells[2].text = 'Şemsiye Fon Türü'
hdr_cells[3].text = 'Getiri'

for i in range (20):
	row_cells = table.add_row().cells
	row_cells[0].text = df1.iloc[i]['Fon Kodu']
	row_cells[1].text = df1.iloc[i]['Fon Adı']
	row_cells[2].text = df1.iloc[i]['Şemsiye Fon Türü']
	row_cells[3].text = str(df1.iloc[i]['Getiri'])
###

##
document.add_page_break()
document.add_heading('Geçen 2 Hafta Bizimkiler - Tablo' , level=1)
document.add_paragraph('Tarih aralığı ' + str(past2) + "----" + str(today), style='Intense Quote')
table = document.add_table(rows=1, cols=4,style='Table Grid')
table.allow_autofit=True

parsing_df1=df1[df1['Fon Kodu'].str.contains("NNF|TCD|AOY|AFS|MAC|AFV|AFA|TGE|GTA|GPU|GUH|GPB|OPH|OPI|ODV|ARM") == True] 

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fon Kodu'
hdr_cells[1].text = 'Fon Adı'
hdr_cells[2].text = 'Şemsiye Fon Türü'
hdr_cells[3].text = 'Getiri'

for i in range (len(parsing_df1)):
	row_cells = table.add_row().cells
	row_cells[0].text = parsing_df1.iloc[i]['Fon Kodu']
	row_cells[1].text = parsing_df1.iloc[i]['Fon Adı']
	row_cells[2].text = parsing_df1.iloc[i]['Şemsiye Fon Türü']
	row_cells[3].text = str(parsing_df1.iloc[i]['Getiri'])
##


document.add_page_break()


## 
# En kötüler geçen 2 hafta

document.add_page_break()

document.add_heading("Geçen 2 Haftanın en Kötüler - Tablo ", level=1)
document.add_paragraph('Tarih aralığı ' + str(past2) + "----" + str(today), style='Intense Quote')


table = document.add_table(rows=1, cols=4,style='Table Grid')
table.allow_autofit=True


hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Fon Kodu'
hdr_cells[1].text = 'Fon Adı'
hdr_cells[2].text = 'Şemsiye Fon Türü'
hdr_cells[3].text = 'Getiri'

for i in range (20):
	row_cells = table.add_row().cells
	row_cells[0].text = df2.iloc[i]['Fon Kodu']
	row_cells[1].text = df2.iloc[i]['Fon Adı']
	row_cells[2].text = df2.iloc[i]['Şemsiye Fon Türü']
	row_cells[3].text = str(df2.iloc[i]['Getiri'])
###



document.add_page_break()

#



document.add_heading('Geçen 2 Haftanın En İyileri - Resim', level=1)
document.add_paragraph('Tarih aralığı ' + str(past2) + "----" + str(today), style='Intense Quote')
p = document.add_paragraph()
r = p.add_run()
r.add_picture(str(current_path) + "/2.png",width=Inches(6.0), height=Inches(7.7))

r = p.add_run()
r.add_picture(str(current_path) + "/2a.png",width=Inches(6.0), height=Inches(7.7))

document.save('Fon_Report_' + str(date_time_file) + '.docx')

